from docx import Document
from docx.shared import Pt,Inches,Cm
import pandas as pd
from datetime import timedelta
import math
from pyecharts.render import make_snapshot
from snapshot_selenium import snapshot
from common.plotbypyecharts import kline_profession
import logging
from logging.handlers import RotatingFileHandler
from common.db_operation import mysql_login
import os


# 日志配置
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
handler = RotatingFileHandler("create_docx.log",
                              encoding="utf-8",
                              maxBytes=10*1024*1024,
                              backupCount=5)
handler.setLevel(logging.INFO)
formatter = logging.Formatter("%(asctime)s - %(name)s - %(levelname)s-%(funcName)s - %(message)s")
handler.setFormatter(formatter)
logger.addHandler(handler)
console = logging.StreamHandler()
console.setFormatter(formatter)
console.setLevel(logging.INFO)
logger.addHandler(console)

# 昨日行情模板

temp_yest = '''
前一交易日（{last_trade_date}）{name}（{code}）开盘{open_price}，最高{high_price}，\
最低{low_price}，收盘{close_price}，相对前收盘点位（{pre_close_price}）涨跌幅为{pct_chg}%，\
其中成交量为{trade_vol}万股，成交金额{trade_amt}亿元。\
'''

# 区间行情模板

temp_period = '''
近{trade_days}个交易日内，{name}（{code}）最高达到{high_price_period}，最低降至{low_price_period}，\
平均点位{mean_price_period}，期间涨跌幅{pct_chg_period}%。\
'''


# 获取最新行情数据
def get_quot_day(df_curr):
    # 最新行情
    last_trade_date = df_curr['日期']
    name = df_curr['简称']
    code = df_curr['代码']
    open_price = round(df_curr['开盘价'],2)
    close_price = round(df_curr['收盘价'],2)
    high_price = round(df_curr['最高价'],2)
    low_price = round(df_curr['最低价'],2)
    pre_close_price = round(df_curr['前收盘价'],2)
    pct_chg = round(df_curr['涨跌幅'],2)
    trade_vol = round(df_curr['成交量'],2)
    trade_amt = round(df_curr['成交额'],2)

    ph_curr = temp_yest.format_map(vars())
    return ph_curr


# 获取区间行情数据
def get_quot_period(df_period):
    picture = kline_profession(df_period)
    code = df_period.values[0][1]
    name = df_period.values[0][2]
    trade_days = int(df_period['收盘价'].describe()['count'])
    low_price_period = round(df_period['收盘价'].describe()['min'],2)
    high_price_period = round(df_period['收盘价'].describe()['max'],2)
    mean_price_period = round(df_period['收盘价'].describe()['mean'],2)
    pct_chg_period = round(calcu_stock_range_period(df_period)*100,2)
    
    ph_period = temp_period.format_map(vars())
    return ph_period, df_period, picture


# 计算股票区间涨跌幅
def calcu_stock_range_period(df_period):
    data = df_period
    stock_range_period = math.exp((data['收盘价']/data['前收盘价']).apply(lambda x: math.log(x)).sum()) - 1
    return stock_range_period

def output_docx(df_period, source_path):
    html_path = source_path +'\\templates\\'
    file_path = source_path +'\\result_data\\'

    # 创建文档
    document = Document()

    # 获取区间行情统计信息
    ph_period, data_dtl, picture = get_quot_period(df_period)

    # 获取最新行情统计信息
    data_curr = data_dtl.iloc[-1]
    ph_curr = get_quot_day(data_curr)

    # 新建段落,写入最新行情统计信息
    paragraph = document.add_paragraph(ph_curr)
    ph_format = paragraph.paragraph_format
    ph_format.space_before = Pt(10)  # 设置段前间距
    ph_format.space_after = Pt(12)   # 设置段后间距
    ph_format.line_spacing = Pt(19)  # 设置行间距

    # 新建段落,写入区间行情统计信息
    paragraph = document.add_paragraph(ph_period)
    ph_format = paragraph.paragraph_format
    ph_format.space_before = Pt(10)  # 设置段前间距
    ph_format.space_after = Pt(12)  # 设置段后间距
    ph_format.line_spacing = Pt(19)  # 设置行间距

    logger.info("行情信息写入完毕...")

    # 新建表格1
    logger.info("正在生成图片...")
    picture.render(html_path+'stock.html')
    make_snapshot(snapshot, picture.render(), file_path+'stock.png')
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    # 插入图片
    run = paragraph.add_run()
    run.add_picture(file_path+'stock.png', width=Inches(6.0))

    # 新建表格2
    logger.info("正在生成表格...")
    data_dtl = data_dtl[['日期', '开盘价', '收盘价', '最低价', '最高价', '成交量', '成交额', '涨跌幅']]
    # data_dtl['日期'] = data_dtl['日期'].apply(lambda x: x.strftime("%Y-%m-%d"))
    data_dtl.rename(columns={'成交量': '成交量/万股', '成交额': '成交额/亿元', '涨跌幅': '涨跌幅/%'}, inplace=True)
    table2 = document.add_table(rows=len(data_dtl.index), cols=len(data_dtl.columns))
    table2.add_row()
    for i in range(len(data_dtl.columns)):
        table2.cell(0, i).text = data_dtl.columns[i]  # 添加表头

    for row in range(1, len(data_dtl.index)+1):
        for col in range(len(data_dtl.columns)):
            table2.cell(row, col).width = 1
            table2.cell(row, col).text = str(data_dtl.iloc[row-1, col])
            table2.cell(row, col).width = Cm(6)
    table2.style = 'Medium Grid 1 Accent 1'
    table2.autofit = True

    # 保存文档
    logger.info("正在生成文档...")
    document.save(file_path + 'doc_test.docx')
    logger.info("文档保存成功...")


if __name__ == "__main__":
    sql = '''
        select  trade_date
                ,index_code
                ,index_name
                ,cast(open_price as dec(20,2)) as open_price
                ,cast(high_price as dec(20,2)) as high_price
                ,cast(low_price as dec(20,2)) as low_price
                ,cast(close_price as dec(20,2)) as close_price
                ,cast(pre_close_price as dec(20,2)) as pre_close_price
                ,cast(trade_vol/1e4 as dec(20,2)) as trade_vol
                ,cast(trade_amt/1e8 as dec(20,2)) as trade_amt
                ,cast(pct_chg as dec(20,2)) as pct_chg
        from    ext_data_stock.index_quotation_info
        where   index_code = 'sh.{0}'
        AND     trade_date BETWEEN '{1}' AND '{2}';
    '''.format('000001', '20200101','20200323')
    try:
        conn = mysql_login()
        df = pd.read_sql(sql=sql, con=conn)
        df['trade_date'] = df['trade_date'].astype('str')
        df.rename(columns={   'trade_date':'日期'
                        ,'index_code':'代码'
                        ,'index_name':'简称'
                        ,'open_price':'开盘价'
                        ,'high_price':'最高价'
                        ,'low_price':'最低价'
                        ,'close_price':'收盘价'
                        ,'pre_close_price':'前收盘价'
                        ,'trade_vol':'成交量'
                        ,'trade_amt':'成交额'
                        ,'pct_chg':'涨跌幅'
                        }, inplace=True)
        conn.close()
    except Exception as e:
        print(e)
    output_docx(df, '..\\')

